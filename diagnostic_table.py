import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from styles import set_header_style, set_table_style, set_doc_style, set_orientation

from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json

with open('config.json', encoding='utf8') as f:
    config = json.load(f)

PATH = config['students_file']
TITLE = config['title']
GRADE = config['grade']
COLUMNS = config['columns']
COLUMNS_WIDTH = config['columns_width']
MATTER_PAGE_HEADER = config['matter']
AIM_PAGE_HEADER = config['aim']
HEADER_TEXT = """ESTADO DO PARÁ
PREFEITURA MUNICIPAL DE VITÓRIA DO XINGU
SECRETARIA MUNICIPAL DE EDUCAÇÃO
EMEF DULCINÉIA ALMEIDA DO NASCIMENTO
INEP 15111130"""
DOCUMENT_NAME = f'{MATTER_PAGE_HEADER}_{TITLE}_{GRADE}º ANO.DOCX'



def extract_classrooms(classrooms_series):
    classrooms = classrooms_series.drop_duplicates().copy().tolist()
    classrooms.sort()
    classrooms = list(filter(lambda c: (c[1] == str(GRADE)), classrooms))
    
    return classrooms

def load_students_grade():
    all_student = pd.read_csv(PATH, na_filter=False)
    classrooms_by_grade = extract_classrooms(all_student['TURMA'])

    student_by_grade = {c: list(all_student[all_student['TURMA'] == c]['ALUNOS'])
                        for c in classrooms_by_grade}

    return student_by_grade

def build_table(document, data, cell_image='images/opcoes.png'):
    table = document.add_table(rows=len(data)+1, cols=len(COLUMNS))

    header_cells = table.rows[0].cells
    for cell, text_header in zip(header_cells, COLUMNS):
        cell.text = text_header

    index_cells = table.columns[0].cells[1:]
    for i, cell in enumerate(index_cells):
        cell.text = str(i+1)
    
    students_cells = table.columns[1].cells[1:]
    for cell, student in zip(students_cells, data):
        cell.text = student
    
    for i in range(2, len(COLUMNS)):
        column = table.columns[i]
        for cell in column.cells[1:]:
            cell.paragraphs[0].add_run().add_picture(cell_image, width=Cm(3.0))
    set_table_style(table, COLUMNS_WIDTH)

def add_aditional_text(doc, classroom, image='images/opcoes.png'):
    doc.add_paragraph()
    par = doc.add_paragraph(TITLE)
    r = par.runs[-1]
    font = r.font
    font.bold = True
    font.size = Pt(12)
    par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # par = doc.add_paragraph('DATA:____________')
    # par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # par = doc.add_paragraph(f'PROFESSOR (A):__________________________________________\tTURMA: {classroom}')
    par = doc.add_paragraph(f'TURMA: {classroom}')
    par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    par.paragraph_format.line_spacing = Pt(10)
    par.paragraph_format.space_after = Pt(2)
    par = doc.add_paragraph(f'COMPONENTE CURRICULAR: {MATTER_PAGE_HEADER}')
    par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    par.paragraph_format.line_spacing = Pt(10)
    par.paragraph_format.space_after = Pt(2)
    par = doc.add_paragraph(f'META: {AIM_PAGE_HEADER}')
    par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    par.paragraph_format.line_spacing = Pt(10)
    par.paragraph_format.space_after = Pt(2)
    par = doc.add_paragraph(f'MARQUE DE ACORDO COM:')
    par.add_run().add_picture(image, width=Cm(4.0))
    par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    par.paragraph_format.line_spacing = Pt(10)
    par.paragraph_format.space_after = Pt(2)

def build_intervention_page(doc):
    doc.add_paragraph()
    par = doc.add_paragraph("ESTRATÉGIAS (INTERVENÇÕES PARA ALCANÇAR A META)")
    r = par.runs[-1]
    font = r.font
    font.bold = True
    font.size = Pt(12)
    par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    line = 110
    par = doc.add_paragraph('_' * line * 19)
    par.paragraph_format.line_spacing = Pt(20)
def build_page(doc, student, classroom):
    add_aditional_text(doc, classroom)
    build_table(doc, student)
    doc.add_page_break()
    build_intervention_page(doc)
      
def build_header(doc, left_img = './images/bandeira-municipio.png', rigth_img = 'images/dulcineia.png'):
    header = doc.sections[0].header

    table = header.add_table(rows=1, cols=3, width=Cm(20.0))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for r in table.rows:
        for c in r.cells:
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    header_text_cell = table.rows[0].cells[1]
    header_text_cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_header_text = header_text_cell.paragraphs[0].add_run(HEADER_TEXT)
    run_header_text.font.size = Pt(8)

    set_header_style(table)

    #set images
    table.cell(0,0).paragraphs[0].add_run().add_picture(left_img, width=Cm(1.77))
    table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0,2).paragraphs[0].add_run().add_picture(rigth_img, width=Cm(1.77))
    table.cell(0,2).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
def build_doc(students_classroom):
    doc = Document()
    
    set_orientation(doc)
    set_doc_style(doc)

    # build_header(doc)
    for i, classroom in enumerate(students_classroom):
        build_page(doc, students_classroom[classroom], classroom)
        if(i < len(students_classroom) - 1):
            doc.add_page_break()

    doc.save(DOCUMENT_NAME)
    
students_classroom = load_students_grade()
build_doc(students_classroom)